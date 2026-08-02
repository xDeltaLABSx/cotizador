import streamlit as st
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import io
import requests

# --- CONFIGURACIÓN DE WORD ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generar_documento_word(data):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # 1. Encabezado
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_company = p_header.add_run("DELTA Land Aerial Building Surveyors LABS\n")
    run_company.bold = True
    run_company.font.size = Pt(12)
    run_company.font.color.rgb = RGBColor(26, 82, 118)
    
    run_sub = p_header.add_run("Servicios Profesionales de Topografía y Geodesia\nCotización de Servicios")
    run_sub.font.size = Pt(9)
    run_sub.font.color.rgb = RGBColor(120, 144, 156)
    
    doc.add_paragraph()
    
    # 2. Fecha
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"Fecha: {data['fecha']}").bold = True
    
    # 3 & 4. Cliente y Empresa
    p_client = doc.add_paragraph()
    p_client.add_run("A quien corresponda:\n").bold = True
    p_client.add_run(f"At'n: {data['cliente']}\n")
    p_client.add_run(f"Empresa: {data['empresa']}").bold = True
    
    doc.add_paragraph()
    
    # 5. Objetivo
    doc.add_heading(level=2).add_run("1. Objetivo del Proyecto").font.color.rgb = RGBColor(26, 82, 118)
    doc.add_paragraph(data['objetivo'])
    
    # 6. Metodología
    doc.add_heading(level=2).add_run("2. Metodología de Trabajo").font.color.rgb = RGBColor(26, 82, 118)
    doc.add_paragraph(data['metodologia'])
    
    # 7. Equipo
    doc.add_heading(level=2).add_run("3. Equipamiento a Utilizar").font.color.rgb = RGBColor(26, 82, 118)
    doc.add_paragraph(data['equipo'])
    
    # 8. Precio (Tabla)
    doc.add_heading(level=2).add_run("4. Propuesta Económica").font.color.rgb = RGBColor(26, 82, 118)
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Descripción del Servicio", "Cant.", "Importe Total"]
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = title
        set_cell_background(cell, "1A5276")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    row_cells = table.rows[1].cells
    row_cells[0].text = data['descripcion_servicio']
    row_cells[1].text = data['cantidad']
    row_cells[2].text = data['precio']
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # 9. Cláusulas
    doc.add_heading(level=2).add_run("5. Términos y Cláusulas").font.color.rgb = RGBColor(26, 82, 118)
    doc.add_paragraph(data['clausulas'])
    
    # 10. Saludo final y Firma
    doc.add_paragraph(data['saludo'])
    
    p_sign = doc.add_paragraph()
    p_sign.add_run("Atentamente,\n\n\n").font.size = Pt(10)
    run_name = p_sign.add_run("Ing. Fernando Cristofer Cárdenas Martínez\n")
    run_name.bold = True
    run_name.font.color.rgb = RGBColor(26, 82, 118)
    p_sign.add_run("Representante Legal y Director Técnico\nDELTA Land Aerial Building Surveyors LABS")
    
    # 11. Pie de página
    for section in doc.sections:
        footer = section.footer
        p_ft = footer.paragraphs[0]
        p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ft = p_ft.add_run("DELTA Land Aerial Building Surveyors LABS — Servicios Profesionales de Topografía y Geodesia")
        run_ft.font.size = Pt(8.5)
        run_ft.font.color.rgb = RGBColor(149, 165, 166)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ STREAMLIT ---
st.set_page_config(page_title="Cotizador - DELTA LABS", page_icon="📐", layout="centered")

st.title("📐 DELTA Land Aerial Building Surveyors LABS")
st.subheader("Generador de Cotizaciones con Base de Datos en Google Drive")

# Inicializar memoria de sesión para el archivo
if "file_ready" not in st.session_state:
    st.session_state.file_ready = False
if "file_data" not in st.session_state:
    st.session_state.file_data = None
if "empresa_nombre" not in st.session_state:
    st.session_state.empresa_nombre = "Empresa"

with st.form("form_cotizacion"):
    fecha = st.text_input("Fecha", value=datetime.now().strftime("%d de %B de %Y"))
    
    st.markdown("#### Destinatario")
    cliente = st.text_input("A quién va dirigida (At'n)", value="Ing. Juan Pérez / Departamento Técnico")
    empresa = st.text_input("Empresa", value="Constructora e Inmobiliaria del Centro, S.A. de C.V.")
    
    st.markdown("#### Detalles Técnicos")
    objetivo = st.text_area("Objetivo", value="Realizar el levantamiento topográfico planialtimétrico de detalle y establecimiento de vértices de control geodésico para la delimitación, trazo y análisis altimétrico del área de estudio requerida.")
    metodologia = st.text_area("Metodología", value="• Reconocimiento de campo y enlace con la Red Geodésica Nacional.\n• Monumentación de bancos de nivel y vértices principales con estacas de acero.\n• Levantamiento con tecnología GPS RTK en doble frecuencia y Estación Total.\n• Procesamiento de datos en gabinete y generación de entregables compatibles con Trimble Coordinate Manager.")
    equipo = st.text_area("Equipo a utilizar", value="• Sistema GNSS RTK de Doble Frecuencia (Base y Rover).\n• Estación Total de alta precisión.\n• Vehículo aéreo no tripulado (Dron) para fotogrametría.\n• Software profesional (AutoCAD, Leica Infinity / QGIS).")
    
    st.markdown("#### Propuesta Económica")
    desc_serv = st.text_input("Descripción del Servicio", value="Levantamiento Topográfico Planialtimétrico y Control Geodésico")
    cantidad = st.text_input("Cantidad", value="1 Lote")
    precio = st.text_input("Importe Total ($ MXN)", value="$35,000.00 MXN")
    
    st.markdown("#### Términos y Cierre")
    clausulas = st.text_area("Cláusulas", value="• Vigencia de la cotización: 15 días.\n• Forma de pago: 50% de anticipo y 50% contra entrega.\n• Tiempo estimado de ejecución: 4 días hábiles.\n• Los precios no incluyen I.V.A.")
    saludo = st.text_area("Saludo final", value="Agradeciendo de antemano su confianza, quedo a su entera disposición para cualquier aclaración o ajuste técnico necesario.")
    
    submitted = st.form_submit_button("Generar Cotización y Enviar a Google Drive")
    
    if submitted:
        data_dict = {
            'fecha': fecha,
            'cliente': cliente,
            'empresa': empresa,
            'objetivo': objetivo,
            'metodologia': metodologia,
            'equipo': equipo,
            'descripcion_servicio': desc_serv,
            'cantidad': cantidad,
            'precio': precio,
            'clausulas': clausulas,
            'saludo': saludo
        }
        
        # Conexión automática con tu Google Drive (Reemplaza con tu URL real)
        webhook_url = "PEGAS_AQUI_TU_URL_DE_GOOGLE_APPS_SCRIPT"
        
        try:
            response = requests.post(webhook_url, json=data_dict)
            if response.status_code == 200:
                st.success("¡Cotización guardada exitosamente en tu Google Drive!")
            else:
                st.warning("Se procesó, pero revisa la conexión con tu Apps Script.")
        except Exception as e:
            st.error(f"Error de conexión con Google Drive: {e}")
        
        # Generar archivo en memoria fuera del flujo estricto del form
        st.session_state.file_data = generar_documento_word(data_dict)
        st.session_state.empresa_nombre = empresa
        st.session_state.file_ready = True

# --- BOTÓN DE DESCARGA FUERA DEL FORMULARIO ---
if st.session_state.file_ready:
    st.markdown("---")
    st.success("¡Documento Word listo para descargar!")
    st.download_button(
        label="📥 Descargar Cotización en Word (.docx)",
        data=st.session_state.file_data,
        file_name=f"Cotizacion_{st.session_state.empresa_nombre.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
