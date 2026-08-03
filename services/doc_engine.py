# services/doc_engine.py
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
from config.settings import COMPANY_INFO, COLORS, obtener_fecha_formal

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

FONT_NAME = "Segoe UI"
BODY_SIZE_PT = 10.5
TITLE_SIZE_PT = 11.0

def _hex_a_rgb(hex_str):
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def _set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def _evitar_salto_fila(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def _aplicar_estilo_parrafo(p, size_pt=BODY_SIZE_PT, bold=False, color_hex=COLORS["TEXT_HEX"], keep_next=False, keep_together=True, space_after=4):
    p.paragraph_format.keep_with_next = keep_next
    p.paragraph_format.keep_together = keep_together
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt) # Fuerza el tamaño de 10.5 pt en cada línea de texto
        run.font.bold = bold
        run.font.color.rgb = _hex_a_rgb(color_hex)

def generar_cotizacion_docx(datos):
    """
    Genera el Word limpiando únicamente el cuerpo del documento y respetando
    al 100% el encabezado e imagen originales de tu plantilla_base.docx.
    """
    ruta_plantilla = os.path.join(BASE_DIR, "assets", "plantilla_base.docx")
    
    if os.path.exists(ruta_plantilla):
        doc = docx.Document(ruta_plantilla)
        # Limpiamos únicamente los párrafos y tablas del cuerpo, SIN TOCAR HEADER/FOOTER
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = docx.Document()

    # Metadatos de Windows
    doc.core_properties.author = datos.get("autor_meta", "DELTA")
    doc.core_properties.title = datos.get("titulo_meta", "")
    doc.core_properties.keywords = datos.get("etiquetas_meta", "")

    # --- 1. FECHA FORMAL (Con el espacio exacto para no tocar el membrete superior) ---
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(12)  # Espacio controlado para el cuerpo
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha_texto = obtener_fecha_formal(datos.get("ciudad"), datos.get("fecha_dt"))
    p_date.add_run(fecha_texto)
    _aplicar_estilo_parrafo(p_date, bold=True, space_after=8)
    
    # --- 2. BLOQUE DE DESTINATARIO ---
    p_client = doc.add_paragraph()
    p_client.paragraph_format.space_after = Pt(10)
    p_client.add_run(f"At'n: {datos.get('cliente_atencion', 'Ingeniero Responsable')}\n")
    if datos.get("cliente_cargo"):
        p_client.add_run(f"{datos.get('cliente_cargo')}\n")
    p_client.add_run(f"{datos.get('cliente_empresa', 'Empresa / Constructora')}\n")
    p_client.add_run("Presente.\n")
    p_client.add_run(f"\nRef: Propuesta de Servicios — {datos.get('nombre_proyecto', 'Levantamiento Topográfico')}")
    _aplicar_estilo_parrafo(p_client, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=8)
    
    # --- 3. SECCIÓN 1: OBJETIVO DEL PROYECTO ---
    h1 = doc.add_heading(level=2)
    h1.add_run("1. Objetivo del Proyecto")
    _aplicar_estilo_parrafo(h1, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    p_obj = doc.add_paragraph(datos.get("objetivo", ""))
    _aplicar_estilo_parrafo(p_obj, space_after=6)
    
    # --- 4. SECCIÓN 2: METODOLOGÍA DE TRABAJO ---
    h2 = doc.add_heading(level=2)
    h2.add_run("2. Metodología y Procedimiento Técnico")
    _aplicar_estilo_parrafo(h2, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    p_met = doc.add_paragraph(datos.get("metodologia", ""))
    _aplicar_estilo_parrafo(p_met, space_after=6)
    
    # --- 5. SECCIÓN 3: INSTRUMENTACIÓN Y EQUIPO ---
    h3 = doc.add_heading(level=2)
    h3.add_run("3. Instrumentación y Equipo Desplegado")
    _aplicar_estilo_parrafo(h3, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    p_eq = doc.add_paragraph(datos.get("equipo", ""))
    _aplicar_estilo_parrafo(p_eq, space_after=6)
    
    # --- 6. SECCIÓN 4: ENTREGABLES ---
    h4 = doc.add_heading(level=2)
    h4.add_run("4. Entregables del Proyecto")
    _aplicar_estilo_parrafo(h4, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    entregables_lista = datos.get("entregables", [])
    for idx, item in enumerate(entregables_lista):
        p_ent = doc.add_paragraph(f"• {item}")
        p_ent.paragraph_format.left_indent = Inches(0.2)
        _aplicar_estilo_parrafo(p_ent, space_after=2)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

  # --- 7. SECCIÓN 5: PROPUESTA ECONÓMICA ---
    h5 = doc.add_heading(level=2)
    h5.add_run("5. Propuesta Económica")
    _aplicar_estilo_parrafo(h5, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=5)
    
    # BLINDAJE: Aseguramos que 'conceptos' siempre sea una lista iterable de diccionarios
    raw_conceptos = datos.get("conceptos_economicos", [])
    if isinstance(raw_conceptos, (int, float)):
        conceptos = [{"desc": datos.get("nombre_proyecto", "Servicio Topográfico"), "cant": "1 Lote", "monto": float(raw_conceptos)}]
    elif isinstance(raw_conceptos, dict):
        conceptos = [raw_conceptos]
    elif isinstance(raw_conceptos, list):
        conceptos = raw_conceptos
    else:
        conceptos = [{"desc": "Servicio General", "cant": "1 Lote", "monto": 0.0}]

    num_filas = len(conceptos) + 2
    tabla = doc.add_table(rows=num_filas, cols=3)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Descripción del Servicio / Concepto", "Cant. / Unidad", "Importe (MXN)"]
    for i, title in enumerate(headers):
        celda = tabla.rows[0].cells[i]
        celda.text = title
        _set_cell_background(celda, COLORS["PRIMARY_HEX"])
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _aplicar_estilo_parrafo(p, size_pt=9.5, bold=True, color_hex="FFFFFF")
    
    _evitar_salto_fila(tabla.rows[0])
        
    total_mxn = 0.0
    for idx, cons in enumerate(conceptos):
        monto = float(cons.get("monto", 0.0))
        total_mxn += monto
        row_cells = tabla.rows[idx + 1].cells
        row_cells[0].text = str(cons.get("desc", "Servicio topográfico"))
        row_cells[1].text = str(cons.get("cant", "1 Lote"))
        row_cells[2].text = f"$ {monto:,.2f}"
        
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if idx % 2 == 1:
            for col_c in row_cells:
                _set_cell_background(col_c, COLORS["ZEBRA_HEX"])
        _evitar_salto_fila(tabla.rows[idx + 1])
                
    # Fila de Total
    celda_tot_lbl = tabla.rows[-1].cells[1]
    celda_tot_lbl.text = "TOTAL (Sin I.V.A.):"
    celda_tot_lbl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _aplicar_estilo_parrafo(celda_tot_lbl.paragraphs[0], size_pt=9.5, bold=True, color_hex=COLORS["PRIMARY_HEX"])
    
    celda_tot_val = tabla.rows[-1].cells[2]
    celda_tot_val.text = f"$ {total_mxn:,.2f}"
    celda_tot_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _aplicar_estilo_parrafo(celda_tot_val.paragraphs[0], size_pt=9.5, bold=True, color_hex=COLORS["PRIMARY_HEX"])
    _evitar_salto_fila(tabla.rows[-1])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # --- 8. SECCIÓN 6: PREMISAS Y EXCLUSIONES ---
    h6 = doc.add_heading(level=2)
    h6.add_run("6. Premisas Técnicas y Exclusiones")
    _aplicar_estilo_parrafo(h6, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    exclusiones_lista = datos.get("exclusiones", [])
    for idx, excl in enumerate(exclusiones_lista):
        p_ex = doc.add_paragraph(f"• {excl}")
        p_ex.paragraph_format.left_indent = Inches(0.2)
        _aplicar_estilo_parrafo(p_ex, color_hex=COLORS["SECONDARY_HEX"], space_after=2)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # --- 9. SECCIÓN 7: CONDICIONES DE TRABAJO ---
    h7 = doc.add_heading(level=2)
    h7.add_run("7. Condiciones de Trabajo y Forma de Pago")
    _aplicar_estilo_parrafo(h7, size_pt=TITLE_SIZE_PT, bold=True, color_hex=COLORS["PRIMARY_HEX"], keep_next=True, space_after=3)
    
    p_cl = doc.add_paragraph(datos.get("clausulas", ""))
    _aplicar_estilo_parrafo(p_cl, space_after=5)
    
    p_sal = doc.add_paragraph(datos.get("saludo_final", ""))
    _aplicar_estilo_parrafo(p_sal, space_after=8)
    
    # --- 10. BLOQUE DE FIRMA Y DATOS BANCARIOS ---
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.keep_together = True
    p_sign.add_run("Atentamente,\n\n\n")
    
    run_name = p_sign.add_run(f"{COMPANY_INFO['LEGAL_REP']}\n")
    run_name.bold = True
    run_name.font.color.rgb = _hex_a_rgb(COLORS["PRIMARY_HEX"])
    
    run_role = p_sign.add_run(f"{COMPANY_INFO['ROLE']}\n{COMPANY_INFO['NAME']}\n\n")
    run_role.font.size = Pt(9.0)
    
    run_bank = p_sign.add_run(f"Datos Bancarios para Anticipo: CLABE {COMPANY_INFO['CLABE_ENDING']} ({COMPANY_INFO['BANK_NAME']})")
    run_bank.font.size = Pt(8.5)
    run_bank.font.color.rgb = _hex_a_rgb(COLORS["SECONDARY_HEX"])
    
    _aplicar_estilo_parrafo(p_sign, space_after=0)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
